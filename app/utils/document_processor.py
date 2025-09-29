import streamlit as st
from docx import Document
import PyPDF2
from io import BytesIO
import pandas as pd
import openpyxl
from typing import Optional, Union, Dict, Any

class DocumentProcessor:
    """Utility class for processing various document types"""

    @staticmethod
    def extract_text_from_docx(file_buffer: BytesIO) -> str:
        """Extract text content from a DOCX file"""
        try:
            doc = Document(file_buffer)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n".join(text_content)

        except Exception as e:
            st.error(f"Error processing DOCX file: {str(e)}")
            return ""

    @staticmethod
    def extract_text_from_pdf(file_buffer: BytesIO) -> str:
        """Extract text content from a PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file_buffer)
            text_content = []

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(text.strip())

            return "\n".join(text_content)

        except Exception as e:
            st.error(f"Error processing PDF file: {str(e)}")
            return ""

    @staticmethod
    def extract_text_from_txt(file_buffer: BytesIO) -> str:
        """Extract text content from a plain text file"""
        try:
            # Reset buffer position to beginning
            file_buffer.seek(0)
            content = file_buffer.read().decode('utf-8')
            st.write(f"🔍 TXT file decoded, content length: {len(content)}")
            return content.strip()

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                file_buffer.seek(0)
                content = file_buffer.read().decode('latin-1')
                st.write(f"🔍 TXT file decoded with latin-1, content length: {len(content)}")
                return content.strip()
            except Exception as e:
                st.error(f"Error processing text file with fallback encoding: {str(e)}")
                return ""
        except Exception as e:
            st.error(f"Error processing text file: {str(e)}")
            return ""

    @staticmethod
    def process_document(uploaded_file) -> Optional[str]:
        """Process uploaded document and extract text content"""
        if uploaded_file is None:
            st.error("❌ No file uploaded")
            return None

        file_extension = uploaded_file.name.lower().split('.')[-1]
        st.write(f"📄 Processing file: {uploaded_file.name}")
        st.write(f"📁 File extension: {file_extension}")
        st.write(f"📊 File size: {uploaded_file.size} bytes")

        try:
            # Create BytesIO object from uploaded file
            file_buffer = BytesIO(uploaded_file.read())
            st.write(f"✅ File read into buffer, size: {len(file_buffer.getvalue())} bytes")

            if file_extension == 'docx':
                content = DocumentProcessor.extract_text_from_docx(file_buffer)
            elif file_extension == 'pdf':
                content = DocumentProcessor.extract_text_from_pdf(file_buffer)
            elif file_extension in ['txt', 'text']:
                content = DocumentProcessor.extract_text_from_txt(file_buffer)
            else:
                st.error(f"❌ Unsupported file format: {file_extension}")
                return None

            if content:
                st.write(f"✅ Text extracted successfully, length: {len(content)} characters")
                with st.expander("🔍 Extracted Content Preview (Debug)"):
                    st.text_area("Content", content[:1000] + "..." if len(content) > 1000 else content, height=200)
                return content
            else:
                st.error("❌ No content extracted from file")
                return None

        except Exception as e:
            st.error(f"❌ Error processing document: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            return None

    @staticmethod
    def load_data_file(uploaded_file) -> Optional[pd.DataFrame]:
        """Load data from CSV or Excel files"""
        if uploaded_file is None:
            return None

        file_extension = uploaded_file.name.lower().split('.')[-1]

        try:
            if file_extension == 'csv':
                df = pd.read_csv(uploaded_file)
            elif file_extension in ['xlsx', 'xls']:
                df = DocumentProcessor._load_excel_with_smart_headers(uploaded_file)
            else:
                st.error(f"Unsupported data file format: {file_extension}")
                return None

            return df

        except Exception as e:
            st.error(f"Error loading data file: {str(e)}")
            return None

    @staticmethod
    def _load_excel_with_smart_headers(uploaded_file) -> Optional[pd.DataFrame]:
        """Load Excel file with smart header detection for HtAG files"""
        try:
            # First, try loading with default header (row 0)
            df_default = pd.read_excel(uploaded_file)

            # Check if we have "Unnamed" columns, which indicates headers might be elsewhere
            unnamed_count = sum(1 for col in df_default.columns if str(col).startswith('Unnamed:'))

            if unnamed_count > df_default.shape[1] * 0.8:  # If >80% columns are unnamed
                st.info("🔍 Detected Excel file with non-standard header layout, searching for proper headers...")

                # Try loading with different header rows (0-5)
                for header_row in range(min(6, len(df_default))):
                    try:
                        df_test = pd.read_excel(uploaded_file, header=header_row)

                        # Skip if this row is empty or still has many unnamed columns
                        if df_test.empty:
                            continue

                        unnamed_test = sum(1 for col in df_test.columns if str(col).startswith('Unnamed:'))

                        # If we significantly reduced unnamed columns, this might be the header row
                        if unnamed_test < df_test.shape[1] * 0.3:  # <30% unnamed columns
                            # Check if this looks like property/suburb data
                            column_names = [str(col).lower() for col in df_test.columns]
                            property_indicators = [
                                'area', 'suburb', 'location', 'price', 'rent', 'yield',
                                'population', 'state', 'region', 'distance', 'growth'
                            ]

                            matches = sum(1 for indicator in property_indicators
                                        for col in column_names if indicator in col)

                            if matches >= 3:  # Found at least 3 property-related columns
                                st.success(f"✅ Found proper headers in row {header_row + 1}")
                                return df_test

                    except Exception as e:
                        continue

                # If no better header found, try to use the first data row as headers
                if len(df_default) > 1:
                    try:
                        # Use the first data row as column names
                        new_columns = df_default.iloc[0].fillna('Unknown').astype(str).tolist()
                        df_fixed = df_default.iloc[1:].copy()
                        df_fixed.columns = new_columns
                        df_fixed = df_fixed.reset_index(drop=True)

                        st.info("🔧 Used first data row as column headers")
                        return df_fixed
                    except Exception as e:
                        st.warning(f"Could not fix headers automatically: {e}")

            return df_default

        except Exception as e:
            st.error(f"Error loading Excel file with smart headers: {str(e)}")
            return None

    @staticmethod
    def validate_suburb_data(df: pd.DataFrame) -> Dict[str, Any]:
        """Validate suburb data structure and return validation results"""
        validation_results = {
            "is_valid": False,
            "row_count": 0,
            "column_count": 0,
            "missing_critical_fields": [],
            "data_quality_issues": [],
            "suggestions": []
        }

        if df is None or df.empty:
            validation_results["data_quality_issues"].append("Dataset is empty")
            return validation_results

        validation_results["row_count"] = len(df)
        validation_results["column_count"] = len(df.columns)

        # Check for critical fields
        critical_fields = [
            "Suburb", "State", "Median Price", "Rental Yield on Houses",
            "Distance (km) to CBD", "Population"
        ]

        missing_fields = []
        for field in critical_fields:
            if field not in df.columns:
                # Try to find similar columns
                similar_cols = [col for col in df.columns if field.lower().replace(" ", "").replace("(", "").replace(")", "") in col.lower().replace(" ", "").replace("(", "").replace(")", "")]
                if not similar_cols:
                    missing_fields.append(field)

        validation_results["missing_critical_fields"] = missing_fields

        # Check data quality
        if df.isnull().sum().sum() > (len(df) * len(df.columns) * 0.3):
            validation_results["data_quality_issues"].append("High percentage of missing values (>30%)")

        # Check for reasonable data ranges (with more lenient criteria)
        if "Median Price" in df.columns:
            price_col = df["Median Price"]
            if price_col.dtype in ['object', 'string']:
                # Check if it's just formatting that needs fixing
                try:
                    # Try to convert a sample to see if it's just formatting
                    sample_val = str(price_col.iloc[0]).replace('$', '').replace(',', '').replace('%', '')
                    float(sample_val)
                    validation_results["suggestions"].append("Price column contains text formatting - will be cleaned automatically")
                except:
                    validation_results["data_quality_issues"].append("Price column contains non-numeric data that cannot be converted")
            elif price_col.min() < 10000 or price_col.max() > 50000000:
                validation_results["suggestions"].append("Some price values may be outside typical ranges")

        # Determine overall validity (more lenient for HtAG processing)
        validation_results["is_valid"] = (
            len(missing_fields) <= 3 and
            len(validation_results["data_quality_issues"]) <= 1
        )

        # Generate suggestions
        if missing_fields:
            validation_results["suggestions"].append(f"Consider adding these fields: {', '.join(missing_fields)}")

        if not validation_results["is_valid"]:
            validation_results["suggestions"].append("Review data source and ensure all critical fields are included")

        return validation_results

    @staticmethod
    def clean_suburb_data(df: pd.DataFrame) -> pd.DataFrame:
        """Clean and standardize suburb data"""
        if df is None or df.empty:
            return df

        df_cleaned = df.copy()

        # Remove completely empty rows
        df_cleaned = df_cleaned.dropna(how='all')

        # Clean suburb names
        if 'Suburb' in df_cleaned.columns:
            df_cleaned['Suburb'] = df_cleaned['Suburb'].astype(str).str.strip().str.title()

        # Clean numeric columns
        numeric_columns = df_cleaned.select_dtypes(include=['number']).columns
        for col in numeric_columns:
            # Remove outliers (values beyond 3 standard deviations)
            mean_val = df_cleaned[col].mean()
            std_val = df_cleaned[col].std()
            if not pd.isna(std_val) and std_val > 0:
                df_cleaned[col] = df_cleaned[col].clip(
                    lower=mean_val - 3*std_val,
                    upper=mean_val + 3*std_val
                )

        # Fill missing values with median for numeric columns
        for col in numeric_columns:
            if df_cleaned[col].isnull().sum() > 0:
                median_val = df_cleaned[col].median()
                df_cleaned[col] = df_cleaned[col].fillna(median_val)

        return df_cleaned